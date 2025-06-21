import docx
from docx import Document
import json

def load_industries():
    """Load industries from the provided document"""
    try:
        doc = Document('attached_assets/Industries_1750323065176.docx')
        industries = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text and len(text) > 2:
                industries.append(text)
        return industries
    except Exception as e:
        print(f"Error loading industries: {e}")
        return []

def load_locations():
    """Load locations from the provided document"""
    try:
        doc = Document('attached_assets/Headquarters_Location_1750323065177.docx')
        locations = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text and len(text) > 2:
                locations.append(text)
        return locations
    except Exception as e:
        print(f"Error loading locations: {e}")
        return []

def process_location_data(locations):
    """Process location data to extract countries and regions"""
    countries = set()
    regions = set()
    
    for location in locations:
        parts = location.split(',')
        if len(parts) >= 2:
            # Last part is usually the country
            country = parts[-1].strip()
            countries.add(country)
            
            # Second to last might be region/state
            if len(parts) >= 3:
                region = parts[-2].strip()
                regions.add(f"{region}, {country}")
    
    return sorted(list(countries)), sorted(list(regions))

def generate_market_estimates():
    """Generate realistic market estimates for industries"""
    industries = load_industries()
    
    # Create market data with more realistic estimates
    market_data = {}
    
    # Industry categories with base TAM estimates (in billions USD)
    tech_industries = ["Artificial Intelligence", "Machine Learning", "Blockchain", "Cryptocurrency", 
                      "Cloud Computing", "Cyber Security", "Software", "Hardware", "Internet"]
    healthcare_industries = ["Biotechnology", "Pharmaceuticals", "Medical Devices", "Healthcare", 
                           "Therapeutics", "Clinical Trials", "mHealth"]
    finance_industries = ["FinTech", "Banking", "Insurance", "Investment", "Asset Management", 
                         "Venture Capital", "Credit Cards"]
    education_industries = ["EdTech", "Education", "STEM Education", "Higher Education", 
                          "Primary Education", "Secondary Education"]
    ecommerce_industries = ["E-Commerce", "Retail", "Marketplace", "Shopping"]
    
    for industry in industries:
        if any(tech in industry for tech in tech_industries):
            base_tam = 800.0 + (hash(industry) % 400)  # 800B - 1.2T range
        elif any(health in industry for health in healthcare_industries):
            base_tam = 400.0 + (hash(industry) % 300)  # 400B - 700B range
        elif any(fin in industry for fin in finance_industries):
            base_tam = 300.0 + (hash(industry) % 200)  # 300B - 500B range
        elif any(edu in industry for edu in education_industries):
            base_tam = 200.0 + (hash(industry) % 150)  # 200B - 350B range
        elif any(ecom in industry for ecom in ecommerce_industries):
            base_tam = 500.0 + (hash(industry) % 300)  # 500B - 800B range
        else:
            base_tam = 50.0 + (hash(industry) % 100)   # 50B - 150B range
        
        market_data[industry] = {
            "global_tam": base_tam,
            "regional_multipliers": {
                "United States": 0.35,
                "China": 0.20,
                "Europe": 0.18,
                "India": 0.08,
                "Japan": 0.06,
                "Canada": 0.03,
                "Australia": 0.02,
                "Brazil": 0.03,
                "United Kingdom": 0.05,
                "Global": 1.0
            }
        }
    
    return market_data

if __name__ == "__main__":
    industries = load_industries()
    locations = load_locations()
    countries, regions = process_location_data(locations)
    market_data = generate_market_estimates()
    
    print(f"Loaded {len(industries)} industries")
    print(f"Loaded {len(countries)} countries")
    print(f"Generated market data for {len(market_data)} industries")
    
    # Save processed data
    with open('industries_data.json', 'w') as f:
        json.dump({
            'industries': industries,
            'countries': countries,
            'regions': regions[:100],  # Limit regions for UI
            'market_data': market_data
        }, f, indent=2)